#!/usr/bin/env python3
"""
Generate BlueLotus V2 Improvements Report as a Word document.
"""

from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
OUTPUT_PATH = Path(r"C:\bluelotus3\research") / f"BlueLotus_V2_Improvements_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

# ── Colour palette — stored as (r, g, b) tuples ─────────────────────────────
DEEP_NAVY    = (0x0D, 0x1B, 0x3E)
GOLD         = (0xD4, 0xAF, 0x37)
MID_BLUE     = (0x1A, 0x3A, 0x6B)
LIGHT_GREY   = (0xF5, 0xF5, 0xF5)
WHITE        = (0xFF, 0xFF, 0xFF)
DARK_TEXT    = (0x1C, 0x1C, 0x1C)
STATUS_GREEN  = (0x2E, 0x7D, 0x32)
STATUS_ORANGE = (0xE6, 0x5C, 0x00)
STATUS_GREY   = (0x55, 0x55, 0x55)
STATUS_BLUE   = (0x15, 0x65, 0xC0)


def rgb(t):
    """Convert (r, g, b) tuple to RGBColor."""
    return RGBColor(t[0], t[1], t[2])


def hex_str(t):
    """Convert (r, g, b) tuple to 6-char uppercase hex string."""
    return f"{t[0]:02X}{t[1]:02X}{t[2]:02X}"


def set_cell_bg(cell, color_tuple):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_str(color_tuple))
    tcPr.append(shd)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), val.get("val", "single"))
            el.set(qn("w:sz"), str(val.get("sz", 4)))
            el.set(qn("w:color"), val.get("color", "000000"))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def add_run_fmt(para, text, bold=False, italic=False, size=11,
                color=None, font="Calibri"):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color) if isinstance(color, tuple) else color
    return run


def add_horizontal_rule(doc, color: RGBColor = GOLD):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(6)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), hex_str(color) if isinstance(color, tuple) else f"{int(color):06X}")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para


def add_section_header(doc, number, title):
    """Dark-navy shaded section banner."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.columns[0].width = Inches(6.5)
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, DEEP_NAVY)
    cell.width = Inches(6.5)
    p = cell.paragraphs[0]
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    add_run_fmt(p, f"  {number}  ", bold=True, size=10, color=GOLD, font="Calibri")
    add_run_fmt(p, title.upper(), bold=True, size=11, color=WHITE, font="Calibri")
    doc.add_paragraph()
    return table


def add_sub_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    add_run_fmt(p, text, bold=True, size=11, color=MID_BLUE, font="Calibri")
    add_horizontal_rule(doc, MID_BLUE)
    return p


def add_body(doc, text, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(space_after)
    add_run_fmt(p, text, size=10.5, color=DARK_TEXT)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.8 + level * 0.5)
    p.paragraph_format.space_after = Pt(3)
    # Bold label if colon present
    if " — " in text:
        label, rest = text.split(" — ", 1)
        add_run_fmt(p, label + " — ", bold=True, size=10, color=DARK_TEXT)
        add_run_fmt(p, rest, size=10, color=DARK_TEXT)
    else:
        add_run_fmt(p, text, size=10, color=DARK_TEXT)
    return p


def add_inline_code(doc, label, value, indent=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(3)
    add_run_fmt(p, label, bold=True, size=10, color=MID_BLUE)
    add_run_fmt(p, value, size=10, color=DARK_TEXT, font="Consolas")
    return p


def add_status_pipeline_table(doc):
    headers = ["Stage", "Status", "Detail"]
    rows = [
        ("Research Forecast Pricing",      "✅  LIVE",       "1,134 rows in ticker_forecasts table"),
        ("Brier Tracking / Resolution",    "⏳  COLLECTING", "First resolutions mature June 9, 2026 (7D horizon)"),
        ("Risk Department Report",         "⏸  BLOCKED",     "Awaiting resolved Brier history"),
        ("CIO Decision Report",            "⏸  BLOCKED",     "Awaiting Risk layer"),
        ("Frontend / Publishing",          "⏸  BLOCKED",     "Awaiting CIO Decision Report"),
        ("Cross-Market Confirmation V2.7", "📋  COMMISSIONED","Work order written; build not started"),
    ]

    status_colors = {
        "✅  LIVE":         STATUS_GREEN,
        "⏳  COLLECTING":   STATUS_ORANGE,
        "⏸  BLOCKED":      STATUS_GREY,
        "📋  COMMISSIONED": STATUS_BLUE,
    }

    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Column widths
    widths = [Inches(2.5), Inches(1.6), Inches(2.4)]
    for i, w in enumerate(widths):
        for row in table.rows:
            row.cells[i].width = w

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        set_cell_bg(hdr.cells[i], DEEP_NAVY)
        p = hdr.cells[i].paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        add_run_fmt(p, h, bold=True, size=10, color=WHITE)

    # Data rows
    for r_idx, (stage, status, detail) in enumerate(rows):
        row = table.rows[r_idx + 1]
        if r_idx % 2 == 1:
            for cell in row.cells:
                set_cell_bg(cell, LIGHT_GREY)

        for c_idx, (cell, val) in enumerate(zip(row.cells, [stage, status, detail])):
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            if c_idx == 1:
                add_run_fmt(p, val, bold=True, size=10, color=status_colors.get(status, DARK_TEXT))
            elif c_idx == 0:
                add_run_fmt(p, val, bold=True, size=10, color=DARK_TEXT)
            else:
                add_run_fmt(p, val, size=10, color=DARK_TEXT)

    doc.add_paragraph()


def add_cover_page(doc):
    # Top spacer
    for _ in range(3):
        doc.add_paragraph()

    # Title banner table
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, DEEP_NAVY)
    cell.width = Inches(6.5)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(6)
    add_run_fmt(p, "BLUELOTUS V2", bold=True, size=28, color=GOLD, font="Calibri Light")

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(6)
    add_run_fmt(p2, "SYSTEM IMPROVEMENTS REPORT", bold=True, size=16, color=WHITE, font="Calibri Light")

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after = Pt(20)
    add_run_fmt(p3, "Sprint Review  ·  Accountability Layer  ·  Roadmap", italic=True, size=11, color=GOLD, font="Calibri")

    doc.add_paragraph()

    # Meta table
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    meta_data = [
        ("Report Date",    datetime.now().strftime("%B %d, %Y")),
        ("Classification", "INTERNAL — CIO DISTRIBUTION"),
        ("System",         "BlueLotus V2 — Deterministic Intelligence Pipeline"),
        ("Status",         "Sprint Closed · V2.7 Commissioned"),
    ]
    for r_idx, (label, val) in enumerate(meta_data):
        row = meta_table.rows[r_idx]
        set_cell_bg(row.cells[0], DEEP_NAVY)
        row.cells[0].width = Inches(2.0)
        row.cells[1].width = Inches(4.5)
        p_l = row.cells[0].paragraphs[0]
        p_l.paragraph_format.space_before = Pt(3)
        p_l.paragraph_format.space_after = Pt(3)
        p_l.paragraph_format.left_indent = Cm(0.2)
        add_run_fmt(p_l, label, bold=True, size=10, color=GOLD)
        p_r = row.cells[1].paragraphs[0]
        p_r.paragraph_format.space_before = Pt(3)
        p_r.paragraph_format.space_after = Pt(3)
        p_r.paragraph_format.left_indent = Cm(0.2)
        add_run_fmt(p_r, val, size=10, color=DARK_TEXT)

    doc.add_page_break()


def build_document():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    # ── Cover ────────────────────────────────────────────────────────────────
    add_cover_page(doc)

    # ── Executive Summary ────────────────────────────────────────────────────
    add_section_header(doc, "00", "Executive Summary")

    add_body(doc,
        "This report documents all improvements made to the BlueLotus V2 deterministic intelligence "
        "pipeline during the current sprint. The defining achievement of this sprint is the establishment "
        "of a verifiable scientific accountability layer — something that did not exist in any prior version. "
        "BlueLotus can now, for the first time, measure whether its forecasting models are accurate against "
        "a probabilistic baseline.", space_after=6)

    add_body(doc,
        "Three new subsystems were built and deployed. One major architectural upgrade was commissioned. "
        "The research report pipeline was extended to professional multi-format output. A formal institutional "
        "gap analysis was documented. All prior sprint bug fixes remain in production.", space_after=10)

    add_sub_heading(doc, "Sprint Outcome at a Glance")
    for item in [
        "Brier Score Superforecast Engine — BUILT AND LIVE (1,134 forecast rows in database)",
        "Forecast Resolution Tracker — BUILT AND COLLECTING (first resolutions June 9, 2026)",
        "Method Comparison Report Generator — BUILT AND COLLECTING",
        "Institutional Quant Level Requirements — FORMALLY DOCUMENTED",
        "Cross-Market Confirmation Layer V2.7 — COMMISSIONED, NOT YET BUILT",
        "R6 Research Report Generator — EXCEL + WORD EXPORT ADDED",
        "All prior DEFECT-01 through DEFECT-08 patches — CONFIRMED IN PRODUCTION",
    ]:
        add_bullet(doc, item)

    doc.add_paragraph()

    # ── Section I ────────────────────────────────────────────────────────────
    add_section_header(doc, "I", "Research Forecast Brier Engine — Completed")

    add_body(doc,
        "Work Order BL-RD-WO-20260602-001, commissioned June 2, 2026, is now closed. This is the single most "
        "significant structural improvement in this sprint. Prior to this work order, BlueLotus generated "
        "recommendations with no measurable accuracy layer — it was an intelligent snapshot tool with no feedback "
        "loop. That gap has been closed.")

    add_sub_heading(doc, "bluelotus_superforecast_engine.py")
    add_body(doc,
        "A full probabilistic forecast engine generating predictions for every ticker in the 200-stock universe "
        "across five time horizons using two competing methods. Output is persisted to the ticker_forecasts MySQL "
        "table. Current database state: 1,134 forecast rows.")

    add_inline_code(doc, "Horizons:  ", "7D  ·  14D  ·  30D  ·  60D  ·  90D")
    add_inline_code(doc, "Methods:   ", "BLUELOTUS_CONSERVATIVE  ·  ANALYST_CONSENSUS")
    add_inline_code(doc, "Universe:  ", "200 tickers")

    add_sub_heading(doc, "BLUELOTUS_CONSERVATIVE Method")
    add_body(doc,
        "Internally-derived valuation using EPS × Sector P/E as the base price, with multiplicative "
        "correction factors applied in sequence:")
    for item in [
        "macro_adjustment()         — macro regime and VIX environment factor",
        "strategic_adjustment()     — catalyst presence, earnings proximity",
        "flow_adjustment()          — institutional flow signals",
        "event_adjustment()         — scheduled events, news momentum",
        "data_quality_penalty()     — confidence discount for incomplete data",
        "Thorp Margin of Safety     — minimum 10% discount applied to final price target",
        "Probability ceiling        — clamped at 0.72 (no overconfidence permitted)",
    ]:
        add_bullet(doc, item)

    add_sub_heading(doc, "ANALYST_CONSENSUS Method")
    add_body(doc,
        "Uses the Moomoo analyst average target price as the 90D anchor. Shorter horizons are derived "
        "by scaling inward using sqrt(weight) probability decay. Probability ceiling clamped at 0.76.")
    add_inline_code(doc, "Horizon weights:  ", "{7D: 0.15,  14D: 0.25,  30D: 0.45,  60D: 0.72,  90D: 1.00}")

    add_sub_heading(doc, "forecast_resolution_tracker.py")
    add_body(doc,
        "Reads due forecasts from ticker_forecasts when a horizon has matured, compares predicted price "
        "against the live price from dataset_raw.json, and computes the following accuracy metrics:")
    for item in [
        "Brier Score — BS = (probability − outcome)²  [range: 0.00 to 1.00, lower is better]",
        "Absolute Error — |predicted_price − actual_price|",
        "Percentage Error — absolute_error / actual_price",
        "Directional Accuracy — whether expected_return sign matched actual_return sign",
        "Actual Return % — (actual_price − current_price) / current_price × 100",
    ]:
        add_bullet(doc, item)
    add_body(doc, "Current status: COLLECTING — earliest 7D resolution date is June 9, 2026.", space_after=4)

    add_sub_heading(doc, "forecast_method_comparison.py")
    add_body(doc,
        "Generates the Brier accountability report by querying forecast_resolutions grouped by "
        "prediction_method, horizon_days, and ticker. Produces two outputs: "
        "research_forecast_accuracy_report.txt and forecast_method_comparison_latest.json.")
    add_body(doc,
        "Embedded doctrine printed in every report output:", space_after=2)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.2)
    p.paragraph_format.right_indent = Cm(1.2)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for side in ("left",):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single"); el.set(qn("w:sz"), "12")
        el.set(qn("w:space"), "8")
        el.set(qn("w:color"), hex_str(GOLD))
        pBdr.append(el)
    pPr.append(pBdr)
    add_run_fmt(p, '"Do not claim forecast skill until resolved Brier history exists."',
                italic=True, size=11, color=MID_BLUE)

    add_sub_heading(doc, "Why This Matters")
    add_body(doc,
        "This transforms BlueLotus from a subjective intelligence tool into a verifiable forecasting system. "
        "Once resolutions begin maturing from June 9, the system will produce empirical answers to questions "
        "that were previously unanswerable:")
    for item in [
        "Which forecast method performs better — BLUELOTUS_CONSERVATIVE or ANALYST_CONSENSUS?",
        "At which horizons is each method most accurate?",
        "Which tickers are most predictable vs least predictable?",
        "Does probability calibration improve or degrade over time?",
    ]:
        add_bullet(doc, item)

    doc.add_paragraph()

    # ── Section II ───────────────────────────────────────────────────────────
    add_section_header(doc, "II", "Institutional Quant Level Requirements — Documented")

    add_body(doc,
        "institutional_quant_level_requirements.md formally maps the full gap between BlueLotus's current "
        "prototype state and institutional quant grade. This document makes gaps machine-readable rather than "
        "implicit, and provides the project with a formal destination.")

    add_sub_heading(doc, "The Core Transformation Required")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.2)
    p.paragraph_format.right_indent = Cm(1.2)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    el = OxmlElement("w:left")
    el.set(qn("w:val"), "single"); el.set(qn("w:sz"), "12")
    el.set(qn("w:space"), "8")
    el.set(qn("w:color"), hex_str(GOLD))
    pBdr.append(el)
    pPr.append(pBdr)
    add_run_fmt(p, "Current:  ", bold=True, size=10, color=DARK_TEXT)
    add_run_fmt(p, "Market intelligence snapshot  →  Research report\n", italic=True, size=10, color=DARK_TEXT)
    add_run_fmt(p, "Required: ", bold=True, size=10, color=DARK_TEXT)
    add_run_fmt(p, "Point-in-time data lake  →  Validated features  →  Tested signals  →  "
                   "Risk-aware portfolio construction  →  Execution  →  Attribution  →  Monitoring  →  Audit",
                italic=True, size=10, color=DARK_TEXT)

    add_sub_heading(doc, "12 Required Capability Areas")
    capabilities = [
        ("1",  "Point-in-Time Historical Data Lake",           "Prevents look-ahead bias; enables proper backtesting"),
        ("2",  "Bias Controls",                                "Look-ahead, survivorship, selection bias elimination"),
        ("3",  "Feature Store / Signal Registry",              "Versioned, reusable signal definitions"),
        ("4",  "Backtesting Engine",                           "Walk-forward validation with statistical significance tests"),
        ("5",  "Statistical Validation Layer",                 "p-values, multiple-comparison correction, out-of-sample hold-out"),
        ("6",  "Risk Model",                                   "Factor exposures, VaR, CVaR, drawdown limits"),
        ("7",  "Portfolio Construction and Optimization",      "Mean-variance, risk parity, Kelly position sizing"),
        ("8",  "Execution and Trade Lifecycle",                "Order management, slippage modeling, fill tracking"),
        ("9",  "Data Quality Contracts",                       "Schema enforcement, freshness SLAs, anomaly alerts"),
        ("10", "Model Lifecycle Management",                   "Versioning, champion/challenger regime, deprecation policy"),
        ("11", "Monitoring and Alerts",                        "Live drift detection, Brier score degradation alerts"),
        ("12", "Governance / Compliance / Auditability",       "Decision audit trail, conflict-of-interest controls"),
    ]

    cap_table = doc.add_table(rows=1 + len(capabilities), cols=3)
    cap_table.style = "Table Grid"
    cap_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cap_widths = [Inches(0.35), Inches(2.5), Inches(3.65)]
    for r in cap_table.rows:
        for i, w in enumerate(cap_widths):
            r.cells[i].width = w

    hdr = cap_table.rows[0]
    for i, h in enumerate(["#", "Capability", "Purpose"]):
        set_cell_bg(hdr.cells[i], DEEP_NAVY)
        p = hdr.cells[i].paragraphs[0]
        p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
        add_run_fmt(p, h, bold=True, size=10, color=WHITE)

    for r_idx, (num, cap, purpose) in enumerate(capabilities):
        row = cap_table.rows[r_idx + 1]
        if r_idx % 2 == 1:
            for cell in row.cells:
                set_cell_bg(cell, LIGHT_GREY)
        for c_idx, val in enumerate([num, cap, purpose]):
            p = row.cells[c_idx].paragraphs[0]
            p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
            add_run_fmt(p, val, bold=(c_idx == 1), size=9.5, color=DARK_TEXT)

    doc.add_paragraph()

    add_sub_heading(doc, "5-Phase Roadmap")
    phases = [
        ("Phase 1", "Data Foundation",               "Point-in-time data lake, bias controls, data quality contracts"),
        ("Phase 2", "Research Validation",            "Feature store, backtesting engine, statistical validation layer"),
        ("Phase 3", "Risk and Portfolio Construction","Risk model, portfolio construction and optimization"),
        ("Phase 4", "Execution and Attribution",      "Trade lifecycle, execution modeling, P&L attribution"),
        ("Phase 5", "Governance and Production",      "Model lifecycle management, monitoring, compliance, audit trail"),
    ]
    for phase, name, detail in phases:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.space_after = Pt(4)
        add_run_fmt(p, f"{phase}: {name} — ", bold=True, size=10, color=MID_BLUE)
        add_run_fmt(p, detail, size=10, color=DARK_TEXT)

    doc.add_paragraph()

    # ── Section III ──────────────────────────────────────────────────────────
    add_section_header(doc, "III", "Cross-Market Confirmation Layer — Commissioned (V2.7)")

    add_body(doc,
        "The Cross-Market Confirmation Layer work order defines the next major intelligence upgrade. "
        "The problem it solves was made concrete on June 6, 2026, when AU fell 8.73% and NEM fell 7.96%. "
        "The current system could not definitively answer whether the cause was gold thesis failure, "
        "miner panic liquidation, broad market risk-off, dollar strength, or yield pressure.")

    add_sub_heading(doc, "The Causal Ambiguity Problem")
    add_body(doc, "When AU or NEM falls sharply, the system today cannot answer:")
    for item in [
        "Is gold itself under pressure — or only the miners?",
        "Is the dollar strengthening, creating a mechanical headwind for gold?",
        "Are long yields rising, increasing the opportunity cost of holding gold?",
        "Is this a broad market liquidation event, and gold miners are collateral damage?",
        "Or is the gold thesis itself failing — physical demand collapsing, ETF outflows accelerating?",
    ]:
        add_bullet(doc, item)

    add_sub_heading(doc, "Scope: 60+ New Tickers Across 11 Categories")
    categories = [
        ("Core Index ETFs",      "SPY, QQQ, IWM, RSP"),
        ("Volatility",           "VXX, UVXY"),
        ("Dollar / Currency",    "DXY, UUP, FXE, FXY"),
        ("Treasury / Yield",     "TLT, IEF, SHY, TIP, ^TNX"),
        ("Gold / Precious Metals","GLD, IAU, SLV, GDX, GDXJ, SIL, SILJ"),
        ("Commodity / Inflation","USO, UNG, DBA, DBC, CPER, URA, XME"),
        ("Sector ETFs",          "XLK, XLF, XLE, XLV, XLI, XLY, XLP, XLU, XLB, XLRE, XLC"),
        ("Growth / Value Factors","VUG, VTV, MTUM, QUAL, USMV, VLUE"),
        ("Credit / Liquidity",   "HYG, JNK, LQD, MBB, AGG"),
        ("Global Markets",       "EFA, EEM, FXI, KWEB, EWJ, EWZ, INDA"),
    ]

    cat_table = doc.add_table(rows=1 + len(categories), cols=2)
    cat_table.style = "Table Grid"
    cat_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cat_widths = [Inches(2.3), Inches(4.2)]
    for r in cat_table.rows:
        for i, w in enumerate(cat_widths):
            r.cells[i].width = w

    hdr_row = cat_table.rows[0]
    for i, h in enumerate(["Category", "Tickers"]):
        set_cell_bg(hdr_row.cells[i], DEEP_NAVY)
        p = hdr_row.cells[i].paragraphs[0]
        p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
        add_run_fmt(p, h, bold=True, size=10, color=WHITE)

    for r_idx, (cat, tickers) in enumerate(categories):
        row = cat_table.rows[r_idx + 1]
        if r_idx % 2 == 1:
            for cell in row.cells:
                set_cell_bg(cell, LIGHT_GREY)
        p0 = row.cells[0].paragraphs[0]
        p0.paragraph_format.space_before = Pt(2); p0.paragraph_format.space_after = Pt(2)
        add_run_fmt(p0, cat, bold=True, size=9.5, color=DARK_TEXT)
        p1 = row.cells[1].paragraphs[0]
        p1.paragraph_format.space_before = Pt(2); p1.paragraph_format.space_after = Pt(2)
        add_run_fmt(p1, tickers, size=9.5, color=DARK_TEXT, font="Consolas")

    doc.add_paragraph()

    add_sub_heading(doc, "9 Derived Scores (New Computed Fields in dataset_raw.json)")
    scores = [
        "market_breadth_confirmation_score — SPY / QQQ / IWM relative strength composite",
        "risk_appetite_score               — HYG vs LQD spread, VIX regime, small-cap performance",
        "dollar_pressure_score             — DXY momentum and mean-reversion signal",
        "yield_pressure_score              — TLT rate-of-change vs long-run trend",
        "gold_thesis_confirmation_score    — GLD vs DXY vs ^TNX correlation composite",
        "gold_miner_relative_strength_score— GDX/GDXJ vs GLD relative performance",
        "gold_miner_panic_liquidation_flag — Detects forced selling in miners vs spot gold",
        "sector_etf_rotation_score         — Money flow between defensive and cyclical sectors",
        "credit_stress_score               — HYG/JNK vs LQD spread widening signal",
    ]
    for s in scores:
        add_bullet(doc, s)

    add_sub_heading(doc, "12 Boolean Interpretation Flags")
    flags = [
        ("broad_market_risk_off",         "SPY/QQQ/IWM all declining in correlation"),
        ("tech_led_selloff",              "QQQ underperforming SPY on relative basis"),
        ("small_cap_risk_off",            "IWM significantly underperforming SPY"),
        ("dollar_pressure_active",        "DXY rising above 20D mean — headwind to USD-denominated assets"),
        ("yield_pressure_active",         "^TNX rising — opportunity cost of gold / rate-sensitive stocks"),
        ("credit_stress_active",          "HYG/JNK spreads widening relative to LQD"),
        ("gold_thesis_confirmed",         "GLD holding while miners lag — physical demand intact"),
        ("gold_thesis_tactical_pressure", "GLD declining with miners — thesis under pressure"),
        ("miner_panic_liquidation",       "GDX/GDXJ falling faster than GLD — leverage unwind"),
        ("bank_thesis_confirmed",         "XLF outperforming when yield_pressure_active is true"),
        ("ai_thesis_failure",             "QQQ/XLK underperforming SPY on sustained basis"),
        ("quantum_panic_liquidation",     "Specific quantum-exposure names falling while XLK holds"),
    ]
    for flag, meaning in flags:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.space_after = Pt(3)
        add_run_fmt(p, f"{flag:<38}", bold=True, size=9.5, color=MID_BLUE, font="Consolas")
        add_run_fmt(p, f"  {meaning}", size=9.5, color=DARK_TEXT)

    add_sub_heading(doc, "Intelligence Quality Projection")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(6)
    add_run_fmt(p, "Current system:  ", bold=True, size=11, color=DARK_TEXT)
    add_run_fmt(p, "8.7 / 10", bold=True, size=14, color=STATUS_ORANGE)
    add_run_fmt(p, "     →     After V2.7:  ", bold=True, size=11, color=DARK_TEXT)
    add_run_fmt(p, "9.5 / 10", bold=True, size=14, color=STATUS_GREEN)

    doc.add_paragraph()

    # ── Section IV ───────────────────────────────────────────────────────────
    add_section_header(doc, "IV", "R6 Research Report Generator — Multi-Format Output")

    add_body(doc,
        "The research report generator reached revision R6, adding professional multi-format output capability. "
        "The CIO Operating Letter is now generated in four formats from a single pipeline run.")

    add_sub_heading(doc, "New Output Formats Added in R6")
    formats = [
        (".txt",   "Research archive — human-readable plain text, consistent with V1 format"),
        (".json",  "Machine-readable delivery — consumed by Risk layer, Frontend, and archive"),
        (".xlsx",  "NEW — Excel export with structured tabular data for quantitative review"),
        (".docx",  "NEW — Word document with formatted CIO Letter for distribution"),
    ]
    for ext, desc in formats:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.space_after = Pt(4)
        add_run_fmt(p, f"{ext:<8}", bold=True, size=10.5, color=MID_BLUE, font="Consolas")
        add_run_fmt(p, f"  {desc}", size=10, color=DARK_TEXT)

    add_body(doc,
        "Confirmed live output: BlueLotus_V2_R6_CIO_Word_Report_20260606_221718.docx — "
        "timestamped document confirms the automated Word generation pipeline is operational.")

    doc.add_paragraph()

    # ── Section V ────────────────────────────────────────────────────────────
    add_section_header(doc, "V", "Prior Sprint Bug Fixes — Confirmed in Production")

    add_sub_heading(doc, "DEFECT-01 through DEFECT-08")
    defects = [
        ("DEFECT-01", "ECE analyst % assertion gate",       "Prevented division-by-zero in ECE scoring when analyst data absent"),
        ("DEFECT-02", "Regime label pollution",              "Fixed regime labels bleeding across ticker contexts in the same cycle"),
        ("DEFECT-03", "Portfolio integrity_flag_reason",     "Always populated — previously left null when no flag was triggered"),
        ("DEFECT-04", "ECE evidence tier tagging",           "Each evidence item now tagged with its tier (Strong/Moderate/Weak)"),
        ("DEFECT-05", "Session-aware price capture",         "Price captured at session open vs intraday — distinction now enforced"),
        ("DEFECT-06", "avg_cost field correction",           "avg_cost was computing wrong cost basis for partial fills"),
        ("DEFECT-07", "BUG-007 duplicate signals",           "moomoo_intelligence.py was emitting duplicate signals per cycle"),
        ("DEFECT-08", "Pool exhaustion fix",                 "db.py v2.0 cycle connection pattern replaces per-request connection opens"),
    ]

    bug_table = doc.add_table(rows=1 + len(defects), cols=3)
    bug_table.style = "Table Grid"
    bug_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    bug_widths = [Inches(1.1), Inches(2.1), Inches(3.3)]
    for r in bug_table.rows:
        for i, w in enumerate(bug_widths):
            r.cells[i].width = w

    for i, h in enumerate(["Defect", "Area", "Fix Applied"]):
        set_cell_bg(bug_table.rows[0].cells[i], DEEP_NAVY)
        p = bug_table.rows[0].cells[i].paragraphs[0]
        p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
        add_run_fmt(p, h, bold=True, size=10, color=WHITE)

    for r_idx, (defect_id, area, fix) in enumerate(defects):
        row = bug_table.rows[r_idx + 1]
        if r_idx % 2 == 1:
            for cell in row.cells:
                set_cell_bg(cell, LIGHT_GREY)
        for c_idx, (val, fnt) in enumerate([(defect_id, "Consolas"), (area, "Calibri"), (fix, "Calibri")]):
            p = row.cells[c_idx].paragraphs[0]
            p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
            add_run_fmt(p, val, bold=(c_idx == 0), size=9.5, color=DARK_TEXT, font=fnt)

    doc.add_paragraph()

    add_sub_heading(doc, "Infrastructure Fixes")
    infra = [
        "FRED API timeout (Singapore routing) — replaced with World Bank Open Data API",
        "SpaceNews CDN rate-limiting — replaced with NASA RSS + Space Google News feeds",
    ]
    for item in infra:
        add_bullet(doc, item)

    doc.add_paragraph()

    # ── Section VI ───────────────────────────────────────────────────────────
    add_section_header(doc, "VI", "Pipeline Status — Current State")

    add_body(doc,
        "The following table reflects the full research-to-publishing pipeline as defined in "
        "Work Order BL-RD-WO-20260602-001, with current operational status as of the report date.")
    doc.add_paragraph()
    add_status_pipeline_table(doc)

    add_sub_heading(doc, "Critical Path")
    add_body(doc,
        "The pipeline is strictly sequential. The first unblocking event is the maturation of the first "
        "7D Brier resolution on June 9, 2026. Once resolved history begins accumulating, the Risk Department "
        "layer can begin. Until that date, the Research layer is the only active stage.")

    doc.add_paragraph()

    # ── Closing ───────────────────────────────────────────────────────────────
    add_section_header(doc, "07", "Assessment and Recommended Next Action")

    add_body(doc,
        "The primary improvement this sprint is the establishment of a verifiable scientific accountability "
        "layer that did not exist before. The Brier score framework forces intellectual honesty — the system "
        "itself prints the doctrine that no forecast skill should be claimed until resolved history exists. "
        "1,134 forecasts are now queued. The clock is running.")

    add_body(doc,
        "The secondary improvement is the formalization of the institutional gap. The requirements document "
        "makes explicit what BlueLotus is not yet, which is a prerequisite for building what it needs to become.")

    add_body(doc,
        "The tertiary improvement is the Cross-Market Confirmation work order — which addresses the most "
        "visible intelligence weakness: the inability to distinguish cause from correlation when individual "
        "positions move.")

    add_sub_heading(doc, "Recommended Next Action: Build V2.7 Cross-Market Confirmation Layer")
    add_body(doc,
        "The work order is commissioned and all requirements are specified. The implementation requires "
        "three phases:")
    phases_next = [
        "Phase 1 — Add 60+ ETF tickers to ticker_universe.py and data collection pipeline",
        "Phase 2 — Build cross_market_confirmation engine computing 9 derived scores and 12 flags into dataset_raw.json",
        "Phase 3 — Modify research_report_generator.py to include the Cross-Market Confirmation section in the CIO Letter",
    ]
    for p_item in phases_next:
        add_bullet(doc, p_item)

    add_body(doc,
        "V2.7 can be built in parallel with Brier collection — it does not depend on the resolution pipeline. "
        "There is no reason to wait.", space_after=12)

    add_horizontal_rule(doc, GOLD)

    p_end = doc.add_paragraph()
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_end.paragraph_format.space_before = Pt(6)
    add_run_fmt(p_end, "BLUELOTUS V2  ·  INTERNAL RESEARCH  ·  NOT FOR EXTERNAL DISTRIBUTION",
                italic=True, size=9, color=STATUS_GREY)

    # Save
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    print(f"Saved: {OUTPUT_PATH}")
    return str(OUTPUT_PATH)


if __name__ == "__main__":
    path = build_document()
    print(f"Done → {path}")

