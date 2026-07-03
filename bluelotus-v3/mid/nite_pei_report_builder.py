"""
BlueLotus V3 â€” NITE-PEI Report Builder
=======================================
Generates TXT, Word (.docx), and Excel (.xlsx) report sections from the
nite_pei{} canonical block.

Called by bluelotus_publisher.py after NITE-PEI block is loaded from the
latest v3 cycle folder.

GOVERNANCE: Deterministic only. No LLM. No order generation.
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List
from zoneinfo import ZoneInfo

import json


# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------

_REPORTS_DIR = Path(__file__).resolve().parents[1] / "reports"
NITE_PEI_TXT_PATH  = _REPORTS_DIR / "bluelotus_v3_nite_pei_report.txt"
NITE_PEI_DOCX_PATH = _REPORTS_DIR / "bluelotus_v3_nite_pei_report.docx"
NITE_PEI_XLSX_PATH = _REPORTS_DIR / "bluelotus_v3_nite_pei_report.xlsx"


def _sgt_now() -> str:
    return datetime.now(ZoneInfo("Asia/Singapore")).strftime("%Y-%m-%d %H:%M SGT")


# ---------------------------------------------------------------------------
# TXT section builder
# ---------------------------------------------------------------------------

def build_nite_pei_txt_section(block: Dict[str, Any]) -> str:
    """Render NITE-PEI block as a plain-text report section."""
    if not block:
        return ""

    lines: List[str] = []
    sep = "=" * 78

    lines += [
        "",
        sep,
        "NITE-PEI THESIS ENGINE â€” CIO ADVISORY",
        "News Impact & Thesis Engine for Prospective Event Intelligence",
        sep,
        f"Generated : {block.get('generated_at_sgt', _sgt_now())}",
        f"Schema    : {block.get('schema_version', 'bluelotus_v3_nite_pei_v1.0')}",
        f"Execution : MANUAL_EXECUTION_REQUIRED | LLM_ORDER_GENERATION=FALSE",
        "",
    ]

    # â”€â”€ CKRI â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    ckri = block.get("ckri", 0.0)
    zone = block.get("ckri_zone", "UNKNOWN")
    detail = block.get("ckri_detail", {})
    lines += [
        "COMPOSITE KILL RISK INDEX (CKRI)",
        "-" * 40,
        f"  CKRI Score : {ckri:.4f}",
        f"  Zone       : {zone}",
        f"  Weighted sum       : {detail.get('weighted_sum', 0.0):.4f}",
        f"  Correlation penalty: {detail.get('correlation_penalty_applied', 0.0):.4f}",
        f"  Total weight       : {detail.get('total_weight', 0.0):.4f}",
        "",
    ]

    zone_guidance = {
        "CLEAR":    "Normal operations â€” no de-risk action required.",
        "WATCH":    "Reduce new adds. Monitor kill conditions.",
        "ELEVATED": "Freeze new adds. CIO review required.",
        "HIGH":     "RISK_REVIEW_REQUIRED â€” consider equity de-risk.",
        "CRITICAL": "RAISE_CASH_REVIEW â€” significant de-risk advisory.",
    }
    lines.append(f"  Zone Guidance: {zone_guidance.get(zone, 'CIO review required.')}")
    lines.append("")

    # Kill breakdown
    breakdown = detail.get("kill_breakdown", [])
    if breakdown:
        lines.append("  Kill Condition Breakdown:")
        lines.append(f"  {'Thesis':<36} {'Kill ID':<20} {'Weight':>7} {'P_kill':>8} {'State':<12}")
        lines.append("  " + "-" * 88)
        for kb in breakdown:
            lines.append(
                f"  {str(kb.get('thesis_id','')):<36} "
                f"{str(kb.get('kill_id','')):<20} "
                f"{kb.get('kill_weight', 0):>7.3f} "
                f"{kb.get('P_kill', 0):>8.4f} "
                f"{kb.get('current_state',''):.<12}"
            )
        lines.append("")

    # â”€â”€ Bayesian Formula Reference â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    lines += [
        "BAYESIAN UPDATE FORMULA (applied per event per thesis)",
        "-" * 40,
        "  Step 1 : prior_odds    = P_prior / (1 - P_prior)",
        "  Step 2 : LR_adjusted   = 1.0 + (LR_table[event_class][thesis_type] - 1.0) x (1 - noise_discount)",
        "           noise_discount pulls evidence toward neutral LR=1.0; T1=0.00  T2=0.10  T3=0.25  T4=0.50",
        "  Step 3 : post_odds     = prior_odds Ã— LR_adjusted",
        "  Step 4 : P_posterior   = post_odds / (1 + post_odds)",
        "  Step 5 : clamp to [0.05, 0.95]  (prevents probability from reaching 0 or 1)",
        "  Multi-event: posterior_N becomes prior_(N+1)  (sequential compounding)",
        "",
        "  LR > 1.0 = event is evidence FOR the thesis (raises P)",
        "  LR < 1.0 = event is evidence AGAINST the thesis (lowers P)",
        "  LR = 1.0 = event carries no information (no update)",
        "",
    ]

    # â”€â”€ Thesis Probability Snapshots â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    snapshots = block.get("thesis_probability_snapshots", [])
    if snapshots:
        lines += [
            "THESIS PROBABILITY SNAPSHOTS â€” WITH EVIDENCE",
            "-" * 40,
        ]
        for snap in snapshots:
            tid = snap.get("thesis_id", "UNKNOWN")
            p_prior = snap.get("P_prior", 0.50)
            p   = snap.get("P_posterior", 0.50)
            dp  = snap.get("delta_p", 0.0)
            arrow = "â–²" if dp > 0 else ("â–¼" if dp < 0 else "â€”")
            posture = snap.get("posture", "THESIS_UNCHANGED â€” MONITOR")
            adv = snap.get("advisory_text", "")
            events = snap.get("events_applied", [])
            lines += [
                f"  â”Œâ”€ {tid}",
                f"  â”‚  P: {p_prior:.4f} â†’ {p:.4f}  ({arrow} {abs(dp):.4f})   Posture: {posture}",
                f"  â”‚  Advisory: {adv}",
            ]
            if events:
                lines.append(f"  â”‚  Evidence ({len(events)} event(s) processed):")
                for i, ev in enumerate(events, 1):
                    beq = ev.get("bayesian_equation", {})
                    src_url = ev.get("source_url", "")
                    pub_at  = ev.get("published_at", "")
                    dset_key = ev.get("dataset_key", "")
                    lines += [
                        f"  â”‚   [{i}] {ev.get('event_class','?')}  keyword: '{ev.get('matched_keyword','?')}'  tier: T{ev.get('source_tier','?')}",
                        f"  â”‚       Headline    : {ev.get('raw_headline','')}",
                        f"  â”‚       Source      : {ev.get('source','')}",
                        f"  â”‚       Published   : {pub_at}",
                        f"  â”‚       Dataset Key : {dset_key}",
                        f"  â”‚       Source URL  : {src_url}" if src_url else "  â”‚       Source URL  : [NO SOURCE URL â€” ACCOUNTABILITY BREACH]",
                        f"  â”‚       â”€â”€ Bayesian Equation â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€",
                        f"  â”‚       {beq.get('step_1_prior_odds','')}",
                        f"  â”‚       {beq.get('step_2_lr_adjustment','')}",
                        f"  â”‚       {beq.get('step_3_posterior_odds','')}",
                        f"  â”‚       {beq.get('step_4_posterior_prob','')}",
                        f"  â”‚       {beq.get('step_5_clamp','')}",
                        f"  â”‚       Î” P this step: {ev.get('delta_p_step',0):+.4f}",
                    ]
            else:
                lines.append("  â”‚  Evidence: No matching events this cycle (P unchanged from prior)")
            lines += [
                f"  â”‚  Kill States: " + "  ".join(
                    f"{k}: {v.get('state','?')} (P={v.get('P_kill',0):.3f})"
                    for k, v in snap.get("kill_state_snapshot", {}).items()
                ) if snap.get("kill_state_snapshot") else "  â”‚  Kill States: (none)",
                "  â””" + "â”€" * 70,
                "",
            ]

    # â”€â”€ Kelly Advisories â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    kelly_advisories = block.get("kelly_advisories", [])
    if kelly_advisories:
        lines += [
            "KELLY-NITE POSITION SIZING ADVISORIES",
            "-" * 40,
            f"  {'Thesis':<36} {'f*_kelly':>9} {'Coherence':>10} {'Target USD':>12} {'Delta USD':>12} {'Action'}",
            "  " + "-" * 100,
        ]
        for k in kelly_advisories:
            lines.append(
                f"  {str(k.get('thesis_id','')):<36} "
                f"{k.get('f_star_kelly', 0):>9.4f} "
                f"{k.get('coherence_score', 0):>10.3f} "
                f"${k.get('target_usd_sleeve', 0):>11,.0f} "
                f"${k.get('delta_usd', 0):>+11,.0f} "
                f"{k.get('advisory_text', '')}"
            )
        lines += [
            "",
            "  All Kelly outputs are advisory only. MANUAL_EXECUTION_REQUIRED.",
            "",
        ]

    # â”€â”€ Contradictions â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    contradictions = block.get("nite_pei_contradictions", [])
    p1_count = block.get("nite_pei_p1_count", 0)
    lines += [
        "NITE-PEI CONTRADICTION REGISTER",
        "-" * 40,
        f"  Total contradictions : {block.get('nite_pei_contradiction_count', 0)}",
        f"  P1 (critical)        : {p1_count}",
        "",
    ]
    if contradictions:
        for c in contradictions:
            lines += [
                f"  [{c.get('severity','?')}] {c.get('contradiction_id','')}",
                f"    Rule     : {c.get('rule','')}",
                f"    Conflict : {c.get('conflict_statement','')}",
                f"    Resolution: {c.get('recommended_resolution_path','')}",
                f"    CIO attention required: {c.get('cio_attention_required', False)}",
                "",
            ]
    else:
        lines.append("  No contradictions detected this cycle.")
        lines.append("")

    lines += [
        sep,
        "END NITE-PEI CIO ADVISORY â€” MANUAL_EXECUTION_REQUIRED",
        sep,
        "",
    ]
    return "\n".join(lines)


def append_nite_pei_report(report: str, nite_pei_block: Dict[str, Any]) -> str:
    """Append NITE-PEI section to the main TXT report string."""
    if not nite_pei_block or "NITE-PEI THESIS ENGINE" in report:
        return report
    section = build_nite_pei_txt_section(nite_pei_block)
    if not section.strip():
        return report
    return report.rstrip() + "\n\n" + section


def save_nite_pei_txt(block: Dict[str, Any]) -> Path:
    """Write standalone NITE-PEI TXT report to reports/."""
    section = build_nite_pei_txt_section(block)
    _REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    NITE_PEI_TXT_PATH.write_text(section, encoding="utf-8")
    return NITE_PEI_TXT_PATH


# ---------------------------------------------------------------------------
# Word (.docx) builder
# ---------------------------------------------------------------------------

def save_nite_pei_docx(block: Dict[str, Any]) -> Path:
    """Write NITE-PEI report as a formatted Word document."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as e:
        raise RuntimeError("python-docx not installed") from e

    doc = Document()

    # â”€â”€ Styles â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    def h1(text: str) -> None:
        p = doc.add_heading(text, level=1)
        p.runs[0].font.color.rgb = RGBColor(0x00, 0x4E, 0x98)

    def h2(text: str) -> None:
        p = doc.add_heading(text, level=2)
        p.runs[0].font.color.rgb = RGBColor(0x1A, 0x56, 0x76)

    def para(text: str, bold: bool = False, italic: bool = False) -> None:
        p = doc.add_paragraph(text)
        if bold or italic:
            for run in p.runs:
                run.bold = bold
                run.italic = italic

    def kv(key: str, val: Any) -> None:
        p = doc.add_paragraph()
        run_k = p.add_run(f"{key}: ")
        run_k.bold = True
        p.add_run(str(val))

    # â”€â”€ Title â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    doc.add_heading("BlueLotus V3 â€” NITE-PEI CIO Advisory Report", level=0)
    kv("Generated", block.get("generated_at_sgt", _sgt_now()))
    kv("Schema", block.get("schema_version", "bluelotus_v3_nite_pei_v1.0"))
    kv("Governance", "MANUAL_EXECUTION_REQUIRED | LLM_ORDER_GENERATION=FALSE | ORDER_ROUTING=FALSE")
    doc.add_paragraph()

    # â”€â”€ CKRI â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    h1("1. Composite Kill Risk Index (CKRI)")
    ckri  = block.get("ckri", 0.0)
    zone  = block.get("ckri_zone", "UNKNOWN")
    detail = block.get("ckri_detail", {})

    zone_colors = {
        "CLEAR": RGBColor(0x22, 0xC5, 0x5E),
        "WATCH": RGBColor(0xF5, 0x9E, 0x0B),
        "ELEVATED": RGBColor(0xF9, 0x73, 0x16),
        "HIGH": RGBColor(0xEF, 0x44, 0x44),
        "CRITICAL": RGBColor(0x7F, 0x1D, 0x1D),
    }
    p = doc.add_paragraph()
    p.add_run("CKRI Score: ").bold = True
    r = p.add_run(f"{ckri:.4f}  |  Zone: {zone}")
    r.bold = True
    r.font.color.rgb = zone_colors.get(zone, RGBColor(0, 0, 0))

    kv("Weighted Sum", f"{detail.get('weighted_sum', 0.0):.4f}")
    kv("Correlation Penalty", f"{detail.get('correlation_penalty_applied', 0.0):.4f}")
    kv("Total Weight", f"{detail.get('total_weight', 0.0):.4f}")

    zone_guidance = {
        "CLEAR":    "Normal operations â€” no de-risk action required.",
        "WATCH":    "Reduce new adds. Monitor kill conditions.",
        "ELEVATED": "Freeze new adds. CIO review required.",
        "HIGH":     "RISK_REVIEW_REQUIRED â€” consider equity de-risk.",
        "CRITICAL": "RAISE_CASH_REVIEW â€” significant de-risk advisory.",
    }
    kv("Zone Guidance", zone_guidance.get(zone, "CIO review required."))
    doc.add_paragraph()

    # Kill breakdown table
    breakdown = detail.get("kill_breakdown", [])
    if breakdown:
        h2("Kill Condition Breakdown")
        tbl = doc.add_table(rows=1, cols=5)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        for i, h in enumerate(["Thesis", "Kill ID", "Weight", "P_kill", "State"]):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True
        for kb in breakdown:
            row = tbl.add_row().cells
            row[0].text = str(kb.get("thesis_id", ""))
            row[1].text = str(kb.get("kill_id", ""))
            row[2].text = f"{kb.get('kill_weight', 0):.3f}"
            row[3].text = f"{kb.get('P_kill', 0):.4f}"
            row[4].text = str(kb.get("current_state", ""))
        doc.add_paragraph()

    # â”€â”€ Thesis Snapshots â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    snapshots = block.get("thesis_probability_snapshots", [])
    if snapshots:
        h1("2. Thesis Probability Snapshots")
        for snap in snapshots:
            tid = snap.get("thesis_id", "UNKNOWN")
            p_val = snap.get("P_posterior", 0.50)
            dp    = snap.get("delta_p", 0.0)
            arrow = "â–²" if dp > 0 else ("â–¼" if dp < 0 else "â€”")
            posture = snap.get("posture", "")
            adv = snap.get("advisory_text", "")

            h2(tid)
            kv("P_posterior", f"{p_val:.3f}  ({arrow} {abs(dp):.3f})")
            kv("Posture", posture)
            kv("Advisory", adv)

            # Evidence with source accountability
            events_snap = snap.get("events_applied", [])
            if events_snap:
                p_ev = doc.add_paragraph()
                p_ev.add_run(f"Evidence ({len(events_snap)} event(s) applied):").bold = True
                for ev_i, ev in enumerate(events_snap, 1):
                    beq = ev.get("bayesian_equation", {})
                    src_url = ev.get("source_url", "")
                    pub_at  = ev.get("published_at", "")
                    dset_key = ev.get("dataset_key", "")
                    p_e = doc.add_paragraph(style="List Bullet")
                    p_e.add_run(f"[{ev_i}] {ev.get('event_class','?')} â€” {ev.get('raw_headline','')}").bold = True
                    kv("Source", ev.get("source", ""))
                    kv("Published", pub_at)
                    kv("Dataset Key", dset_key)
                    p_url = doc.add_paragraph()
                    run_label = p_url.add_run("Source URL: ")
                    run_label.bold = True
                    if src_url:
                        run_link = p_url.add_run(src_url)
                        run_link.font.color.rgb = RGBColor(0x00, 0x56, 0xD2)
                    else:
                        run_link = p_url.add_run("[NO SOURCE URL â€” ACCOUNTABILITY BREACH]")
                        run_link.font.color.rgb = RGBColor(0xEF, 0x44, 0x44)
                    p_eq = doc.add_paragraph()
                    p_eq.add_run("Bayesian Equation:").bold = True
                    for step_key in ["step_1_prior_odds", "step_2_lr_adjustment", "step_3_posterior_odds",
                                     "step_4_posterior_prob", "step_5_clamp"]:
                        step_val = beq.get(step_key, "")
                        if step_val:
                            p_s = doc.add_paragraph(f"    {step_val}", style="Normal")
                            p_s.paragraph_format.left_indent = Inches(0.5)
                            for r in p_s.runs:
                                r.font.name = "Courier New"
                                r.font.size = Pt(9)
                    p_dp = doc.add_paragraph(f"    Î” P this step: {ev.get('delta_p_step',0):+.4f}")
                    p_dp.paragraph_format.left_indent = Inches(0.5)
                    for r in p_dp.runs:
                        r.font.name = "Courier New"
                        r.font.size = Pt(9)
                        r.italic = True

            kill_snap = snap.get("kill_state_snapshot", {})
            if kill_snap:
                p_ks = doc.add_paragraph()
                p_ks.add_run("Kill States: ").bold = True
                p_ks.add_run("  |  ".join(f"{k}: {v.get('state','?')} (P={v.get('P_kill',0):.3f})" for k, v in kill_snap.items()))
            doc.add_paragraph()

    # â”€â”€ Kelly Advisories â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    kelly_advisories = block.get("kelly_advisories", [])
    if kelly_advisories:
        h1("3. Kelly-NITE Position Sizing Advisories")
        para("All outputs are advisory only. MANUAL_EXECUTION_REQUIRED.", italic=True)
        tbl = doc.add_table(rows=1, cols=7)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        for i, h in enumerate(["Thesis", "f*_full", "f*_kelly", "Coherence", "Target USD", "Delta USD", "Action"]):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True
        for k in kelly_advisories:
            row = tbl.add_row().cells
            row[0].text = str(k.get("thesis_id", ""))
            row[1].text = f"{k.get('f_star_full', 0):.4f}"
            row[2].text = f"{k.get('f_star_kelly', 0):.4f}"
            row[3].text = f"{k.get('coherence_score', 0):.3f}"
            row[4].text = f"${k.get('target_usd_sleeve', 0):,.0f}"
            row[5].text = f"${k.get('delta_usd', 0):+,.0f}"
            row[6].text = str(k.get("advisory_text", ""))
        doc.add_paragraph()

    # â”€â”€ Contradictions â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    h1("4. NITE-PEI Contradiction Register")
    kv("Total Contradictions", block.get("nite_pei_contradiction_count", 0))
    kv("P1 Critical", block.get("nite_pei_p1_count", 0))
    doc.add_paragraph()

    contradictions = block.get("nite_pei_contradictions", [])
    if contradictions:
        for c in contradictions:
            h2(f"[{c.get('severity','?')}] {c.get('rule','')} â€” {c.get('contradiction_id','')}")
            kv("Conflict", c.get("conflict_statement", ""))
            kv("Resolution", c.get("recommended_resolution_path", ""))
            kv("CIO Attention Required", c.get("cio_attention_required", False))
            doc.add_paragraph()
    else:
        para("No contradictions detected this cycle.", italic=True)

    # â”€â”€ Footer â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    doc.add_paragraph()
    p = doc.add_paragraph("MANUAL_EXECUTION_REQUIRED â€” CIO authority only â€” No order routing")
    p.runs[0].bold = True

    _REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(NITE_PEI_DOCX_PATH))
    return NITE_PEI_DOCX_PATH


# ---------------------------------------------------------------------------
# Excel (.xlsx) builder
# ---------------------------------------------------------------------------

def save_nite_pei_xlsx(block: Dict[str, Any]) -> Path:
    """Write NITE-PEI report as a multi-tab Excel workbook."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
    except ImportError as e:
        raise RuntimeError("openpyxl not installed") from e

    wb = Workbook()

    # â”€â”€ Style helpers â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    HDR_FILL  = PatternFill("solid", fgColor="004E98")
    HDR_FONT  = Font(bold=True, color="FFFFFF", name="Calibri", size=10)
    BOLD_FONT = Font(bold=True, name="Calibri", size=10)
    NORM_FONT = Font(name="Calibri", size=10)
    ZONE_FILL = {
        "CLEAR":    PatternFill("solid", fgColor="D1FAE5"),
        "WATCH":    PatternFill("solid", fgColor="FEF3C7"),
        "ELEVATED": PatternFill("solid", fgColor="FFEDD5"),
        "HIGH":     PatternFill("solid", fgColor="FEE2E2"),
        "CRITICAL": PatternFill("solid", fgColor="7F1D1D"),
    }
    ZONE_FONT = {
        "CRITICAL": Font(bold=True, color="FFFFFF", name="Calibri", size=10),
    }
    thin = Side(style="thin")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    def hdr_row(ws: Any, row: int, headers: List[str]) -> None:
        for col, h in enumerate(headers, 1):
            c = ws.cell(row=row, column=col, value=h)
            c.font = HDR_FONT
            c.fill = HDR_FILL
            c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            c.border = border

    def data_cell(ws: Any, row: int, col: int, value: Any,
                  bold: bool = False, fill: Any = None, num_fmt: str = None) -> Any:
        c = ws.cell(row=row, column=col, value=value)
        c.font = BOLD_FONT if bold else NORM_FONT
        c.alignment = Alignment(vertical="center", wrap_text=True)
        c.border = border
        if fill:
            c.fill = fill
        if num_fmt:
            c.number_format = num_fmt
        return c

    def auto_width(ws: Any, min_w: int = 10, max_w: int = 40) -> None:
        for col_cells in ws.columns:
            max_len = max((len(str(c.value or "")) for c in col_cells), default=0)
            ws.column_dimensions[get_column_letter(col_cells[0].column)].width = min(max(max_len + 2, min_w), max_w)

    # â”€â”€ Tab 1: CKRI Summary â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    ws1 = wb.active
    ws1.title = "CKRI Summary"
    ws1.append(["BlueLotus V3 â€” NITE-PEI CIO Advisory"])
    ws1["A1"].font = Font(bold=True, size=14, name="Calibri", color="004E98")
    ws1.append(["Generated", block.get("generated_at_sgt", _sgt_now())])
    ws1.append(["Schema", block.get("schema_version", "")])
    ws1.append(["Governance", "MANUAL_EXECUTION_REQUIRED | LLM_ORDER_GENERATION=FALSE"])
    ws1.append([])

    zone = block.get("ckri_zone", "UNKNOWN")
    ckri = block.get("ckri", 0.0)
    detail = block.get("ckri_detail", {})
    ws1.append(["CKRI SCORE", f"{ckri:.4f}"])
    ws1["A6"].font = BOLD_FONT
    ws1["B6"].font = Font(bold=True, size=12, name="Calibri")
    ws1.append(["CKRI ZONE", zone])
    ws1["A7"].font = BOLD_FONT
    ws1["B7"].fill = ZONE_FILL.get(zone, PatternFill())
    if zone in ZONE_FONT:
        ws1["B7"].font = ZONE_FONT[zone]
    ws1.append(["Weighted Sum", detail.get("weighted_sum", 0.0)])
    ws1.append(["Correlation Penalty", detail.get("correlation_penalty_applied", 0.0)])
    ws1.append(["Total Weight", detail.get("total_weight", 0.0)])
    ws1.append([])

    breakdown = detail.get("kill_breakdown", [])
    if breakdown:
        ws1.append(["Kill Condition Breakdown"])
        ws1[f"A{ws1.max_row}"].font = BOLD_FONT
        r = ws1.max_row + 1
        hdr_row(ws1, r, ["Thesis", "Kill ID", "Weight", "P_kill", "State", "Contribution"])
        for kb in breakdown:
            r += 1
            state = str(kb.get("current_state", ""))
            fill = {"CONFIRMED": PatternFill("solid", fgColor="FEE2E2"),
                    "TRIGGERED": PatternFill("solid", fgColor="FFEDD5"),
                    "WATCH":     PatternFill("solid", fgColor="FEF3C7"),
                    "INACTIVE":  PatternFill("solid", fgColor="D1FAE5")}.get(state)
            data_cell(ws1, r, 1, kb.get("thesis_id", ""))
            data_cell(ws1, r, 2, kb.get("kill_id", ""))
            data_cell(ws1, r, 3, kb.get("kill_weight", 0), num_fmt="0.000")
            data_cell(ws1, r, 4, kb.get("P_kill", 0), num_fmt="0.0000")
            data_cell(ws1, r, 5, state, fill=fill)
            data_cell(ws1, r, 6, kb.get("contribution", 0), num_fmt="0.0000")
    auto_width(ws1)

    # â”€â”€ Tab 2: Thesis Probabilities â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    ws2 = wb.create_sheet("Thesis Probabilities")
    hdr_row(ws2, 1, ["Thesis ID", "P_prior", "P_posterior", "Delta P", "Posture", "Advisory Text"])
    r2 = 1
    URL_FONT = Font(name="Calibri", size=9, color="0056D2", underline="single")
    BREACH_FONT = Font(name="Calibri", size=9, color="EF4444", bold=True)
    EQ_FONT  = Font(name="Courier New", size=9)
    META_FONT = Font(name="Calibri", size=9, italic=True, color="555555")
    for snap in block.get("thesis_probability_snapshots", []):
        r2 += 1
        dp = snap.get("delta_p", 0.0)
        data_cell(ws2, r2, 1, snap.get("thesis_id", ""), bold=True)
        data_cell(ws2, r2, 2, snap.get("P_prior", 0.50), num_fmt="0.0000")
        data_cell(ws2, r2, 3, snap.get("P_posterior", 0.50), num_fmt="0.0000")
        c = data_cell(ws2, r2, 4, dp, num_fmt="+0.0000;-0.0000;0.0000")
        c.fill = PatternFill("solid", fgColor="D1FAE5") if dp > 0 else (
                 PatternFill("solid", fgColor="FEE2E2") if dp < 0 else PatternFill())
        data_cell(ws2, r2, 5, snap.get("posture", ""))
        data_cell(ws2, r2, 6, snap.get("advisory_text", ""))
        # Expand evidence rows
        events = snap.get("events_applied", [])
        if events:
            r2 += 1
            ws2.merge_cells(start_row=r2, start_column=1, end_row=r2, end_column=6)
            c = ws2.cell(row=r2, column=1, value="  EVIDENCE â€” Bayesian Update Steps (with Source Accountability)")
            c.font = Font(bold=True, italic=True, name="Calibri", size=9, color="1A5676")
            for ev_i, ev in enumerate(events, 1):
                beq = ev.get("bayesian_equation", {})
                src_url  = ev.get("source_url", "")
                pub_at   = ev.get("published_at", "")
                dset_key = ev.get("dataset_key", "")
                # Event header row
                r2 += 1
                ws2.cell(row=r2, column=1, value=f"    [{ev_i}] {ev.get('event_class','?')}  kw: '{ev.get('matched_keyword','?')}'  T{ev.get('source_tier','?')}").font = BOLD_FONT
                ws2.merge_cells(start_row=r2, start_column=2, end_row=r2, end_column=6)
                ws2.cell(row=r2, column=2, value=ev.get("raw_headline", "")).font = NORM_FONT
                # Source + accountability row
                r2 += 1
                ws2.cell(row=r2, column=1, value="      Source:").font = BOLD_FONT
                ws2.merge_cells(start_row=r2, start_column=2, end_row=r2, end_column=3)
                ws2.cell(row=r2, column=2, value=ev.get("source", "")).font = META_FONT
                ws2.cell(row=r2, column=4, value="Published:").font = BOLD_FONT
                ws2.cell(row=r2, column=5, value=pub_at).font = META_FONT
                ws2.cell(row=r2, column=6, value=dset_key).font = META_FONT
                # Source URL row â€” highlighted
                r2 += 1
                ws2.merge_cells(start_row=r2, start_column=1, end_row=r2, end_column=6)
                if src_url:
                    c_url = ws2.cell(row=r2, column=1, value=f"      SOURCE URL: {src_url}")
                    c_url.font = URL_FONT
                    c_url.fill = PatternFill("solid", fgColor="EFF6FF")
                else:
                    c_url = ws2.cell(row=r2, column=1, value="      SOURCE URL: [NO SOURCE URL â€” ACCOUNTABILITY BREACH]")
                    c_url.font = BREACH_FONT
                    c_url.fill = PatternFill("solid", fgColor="FEE2E2")
                # Bayesian equation rows
                for key in ["step_1_prior_odds", "step_2_lr_adjustment",
                            "step_3_posterior_odds", "step_4_posterior_prob", "step_5_clamp"]:
                    step_val = beq.get(key, "")
                    if step_val:
                        r2 += 1
                        ws2.merge_cells(start_row=r2, start_column=2, end_row=r2, end_column=6)
                        ws2.cell(row=r2, column=2, value=f"        {step_val}").font = EQ_FONT
                r2 += 1
                ws2.merge_cells(start_row=r2, start_column=2, end_row=r2, end_column=6)
                ws2.cell(row=r2, column=2, value=f"        Î” P this step: {ev.get('delta_p_step',0):+.4f}").font = Font(name="Calibri", size=9, italic=True)
        r2 += 1  # spacer
    auto_width(ws2)

    # â”€â”€ Tab 3: Kelly Advisories â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    ws3 = wb.create_sheet("Kelly Advisories")
    hdr_row(ws3, 1, [
        "Thesis ID", "Thesis Type", "P_posterior", "Upside %", "f*_full", "Frac Mult",
        "f*_kelly", "Coherence", "H_norm", "Dispersion",
        "NAV Total", "Target USD", "Current USD", "Delta USD", "Action"
    ])
    for i, k in enumerate(block.get("kelly_advisories", []), 2):
        delta = k.get("delta_usd", 0.0)
        data_cell(ws3, i,  1, k.get("thesis_id", ""))
        data_cell(ws3, i,  2, k.get("thesis_type", ""))
        data_cell(ws3, i,  3, k.get("p_posterior_used", 0), num_fmt="0.000")
        data_cell(ws3, i,  4, k.get("analyst_upside_pct", 0), num_fmt="0.0%")
        data_cell(ws3, i,  5, k.get("f_star_full", 0), num_fmt="0.0000")
        data_cell(ws3, i,  6, k.get("fractional_multiplier", 0), num_fmt="0.000")
        data_cell(ws3, i,  7, k.get("f_star_kelly", 0), num_fmt="0.0000")
        data_cell(ws3, i,  8, k.get("coherence_score", 0), num_fmt="0.000")
        data_cell(ws3, i,  9, k.get("h_norm_used", 0), num_fmt="0.000")
        data_cell(ws3, i, 10, k.get("dispersion_used", 0), num_fmt="0.000")
        data_cell(ws3, i, 11, k.get("nav_total", 0), num_fmt='$#,##0')
        data_cell(ws3, i, 12, k.get("target_usd_sleeve", 0), num_fmt='$#,##0')
        data_cell(ws3, i, 13, k.get("current_usd_sleeve", 0), num_fmt='$#,##0')
        c = data_cell(ws3, i, 14, delta, num_fmt='$#,##0;($#,##0)')
        c.fill = PatternFill("solid", fgColor="D1FAE5") if delta > 0 else (
                 PatternFill("solid", fgColor="FEE2E2") if delta < 0 else PatternFill())
        data_cell(ws3, i, 15, k.get("advisory_text", ""))
    auto_width(ws3)

    # â”€â”€ Tab 4: Contradictions â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    ws4 = wb.create_sheet("Contradictions")
    hdr_row(ws4, 1, ["Contradiction ID", "Severity", "Rule", "Domain", "Source A", "Source B", "Conflict Statement", "Resolution Path", "CIO Attention"])
    contradictions = block.get("nite_pei_contradictions", [])
    if contradictions:
        sev_fill = {"P1": PatternFill("solid", fgColor="FEE2E2"), "P3": PatternFill("solid", fgColor="FEF3C7")}
        for i, c in enumerate(contradictions, 2):
            sev = c.get("severity", "")
            fill = sev_fill.get(sev)
            data_cell(ws4, i, 1, c.get("contradiction_id", ""), fill=fill)
            data_cell(ws4, i, 2, sev, bold=True, fill=fill)
            data_cell(ws4, i, 3, c.get("rule", ""), fill=fill)
            data_cell(ws4, i, 4, c.get("domain", ""))
            data_cell(ws4, i, 5, c.get("source_a", ""))
            data_cell(ws4, i, 6, c.get("source_b", ""))
            data_cell(ws4, i, 7, c.get("conflict_statement", ""))
            data_cell(ws4, i, 8, c.get("recommended_resolution_path", ""))
            data_cell(ws4, i, 9, str(c.get("cio_attention_required", False)))
    else:
        ws4.cell(row=2, column=1, value="No contradictions detected this cycle.").font = NORM_FONT
    auto_width(ws4)

    _REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    wb.save(str(NITE_PEI_XLSX_PATH))
    return NITE_PEI_XLSX_PATH
